import os
import re
import io
import json
import uuid
import time
import random
import hashlib
import subprocess
import tempfile
from datetime import datetime
from typing import Dict, Any

from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from openai import AzureOpenAI
from openpyxl import Workbook, load_workbook
from docx import Document

from dotenv import load_dotenv
load_dotenv()

RESUME_SCHEMA = {
    "first_name": None, "middle_name": None, "last_name": None,
    "dob": None, "gender": None, "ssn": None, "dl_number": None, "dl_state": None,
    "marital_status": None, "is_us_citizen": None,
    "street_number": None, "street_name": None, "suite": None,
    "city": None, "state": None, "zip": None,
    "home_phone": None, "work_phone": None, "fax": None,
    "email": None, "other_email": None,
    "emergency_contact_1": None, "relationship_1": None, "phone_1": None,
    "emergency_contact_2": None, "relationship_2": None, "phone_2": None,
    "high_school_name": None, "hs_last_year": None, "hs_graduated": None,
    "college_name": None, "college_last_year": None, "college_graduated": None,
    "trade_school_name": None, "trade_last_year": None, "trade_graduated": None,
    "degrees": None, "academic_awards": None,
    "certifications": [],
    "previous_employments": [],
}

CERTIFICATION_ITEM = {
    "type": None, "issued_date": None, "organization": None, "license_number": None,
}

PREVIOUS_EMPLOYMENT_ITEM = {
    "employer": None, "position": None, "duties": None,
    "start_salary": None, "end_salary": None,
    "start_date": None, "end_date": None,
    "supervisor": None, "contact_person": None, "reason_for_leaving": None,
    "street_number": None, "street_name": None, "suite": None,
    "city": None, "state": None, "zip": None,
    "phone": None, "email": None,
}


def extract_text_with_ocr(file_content: bytes) -> str:
    """Use Azure prebuilt OCR model to extract text from a PDF or image."""
    endpoint = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT")
    key = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_KEY")

    if not endpoint or not key:
        raise ValueError("Azure Document Intelligence credentials are not configured")

    client = DocumentIntelligenceClient(
        endpoint=endpoint,
        credential=AzureKeyCredential(key),
        api_version="2024-11-30",
    )

    poller = client.begin_analyze_document(
        model_id="prebuilt-read",
        body=file_content,
        content_type="application/octet-stream",
    )
    result = poller.result()

    if not result or not hasattr(result, "content"):
        raise ValueError("No text extracted from document")

    return result.content


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from a .docx file directly (no OCR needed)."""
    doc = Document(io.BytesIO(file_content))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def convert_doc_to_docx(file_content: bytes) -> bytes:
    """Convert legacy .doc to .docx via LibreOffice (must be installed)."""
    with tempfile.TemporaryDirectory() as tmpdir:
        doc_path = os.path.join(tmpdir, "in.doc")
        with open(doc_path, "wb") as f:
            f.write(file_content)
        try:
            subprocess.run(
                ["soffice", "--headless", "--convert-to", "docx",
                 "--outdir", tmpdir, doc_path],
                check=True, capture_output=True,
            )
        except FileNotFoundError:
            raise ValueError(
                "LibreOffice ('soffice') not found. Install it to support .doc files, "
                "or ask the user to save the file as .docx."
            )
        except subprocess.CalledProcessError as e:
            raise ValueError(f"Failed to convert .doc to .docx: {e.stderr.decode(errors='ignore')}")

        docx_path = os.path.join(tmpdir, "in.docx")
        with open(docx_path, "rb") as f:
            return f.read()


def extract_with_gpt(ocr_text: str) -> Dict[str, Any]:
    """Use GPT to reformat the extracted text into desired JSON output."""
    openai_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
    openai_key = os.getenv("AZURE_OPENAI_API_KEY")
    deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")
    api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2025-04-01-preview")

    if not openai_endpoint or not openai_key:
        raise ValueError("Azure OpenAI credentials not configured")

    client = AzureOpenAI(
        azure_endpoint=openai_endpoint,
        api_key=openai_key,
        api_version=api_version,
    )

    extraction_id = str(uuid.uuid4())
    random_seed = random.randint(100000, 999999)
    doc_hash = hashlib.sha256(ocr_text.encode()).hexdigest()[:16]

    time.sleep(0.5)

    cache_breaker = random.choice([
        "You must extract data with precision",
        "Your task is to extract resume information",
        "Extract the resume data accurately",
        "Parse the following resume carefully",
    ])

    schema_json = json.dumps(RESUME_SCHEMA, indent=2)
    cert_item_json = json.dumps(CERTIFICATION_ITEM)
    emp_item_json = json.dumps(PREVIOUS_EMPLOYMENT_ITEM)

    prompt = f"""Extract resume data and return JSON matching EXACTLY this schema (use null for missing values, [] for missing lists):

{schema_json}

Each item in `certifications` must follow this shape:
{cert_item_json}

Each item in `previous_employments` must follow this shape:
{emp_item_json}

Rules:
- `is_us_citizen`, `hs_graduated`, `college_graduated`, `trade_graduated` must be true / false / null.
- `hs_last_year` ∈ {{9,10,11,12,null}}, `college_last_year` ∈ {{1,2,3,4,null}}, `trade_last_year` ∈ {{1,2,null}}.
- Dates as MM/DD/YYYY when possible.
- Salaries as numbers when possible (strip $ and commas), else null.
- Return ONLY JSON. No markdown, no commentary.

Resume text:
{ocr_text}
"""

    response = client.chat.completions.create(
        model=deployment_name,
        messages=[
            {
                "role": "system",
                "content": (
                    f"{cache_breaker} SESSION ID: {extraction_id} "
                    f"DOCUMENT HASH: {doc_hash} "
                    "You are a precise document data extractor. Return only valid JSON."
                ),
            },
            {"role": "user", "content": prompt},
        ],
        temperature=0.1,
        max_completion_tokens=8000,
        seed=random_seed,
        user=f"extraction_{extraction_id}",
        response_format={"type": "json_object"},
    )

    result_text = response.choices[0].message.content or ""
    result_text = re.sub(r"^```(?:json)?\s*|\s*```$", "", result_text.strip())

    try:
        result = json.loads(result_text)
    except json.JSONDecodeError:
        cleaned = re.sub(r",\s*([}\]])", r"\1", result_text)
        try:
            result = json.loads(cleaned)
        except json.JSONDecodeError as e:
            raise ValueError(
                f"GPT returned invalid JSON: {e}\n"
                f"Raw response (first 500 chars): {result_text[:500]}"
            )

    for key, default in RESUME_SCHEMA.items():
        if key not in result or (result[key] is None and isinstance(default, list)):
            result[key] = [] if isinstance(default, list) else default

    return result


def flatten_for_excel(data: Dict[str, Any], filename: str) -> Dict[str, Any]:
    """Flatten nested lists into JSON strings so the record fits one row."""
    flat = dict(data)
    flat["certifications"] = json.dumps(data.get("certifications") or [], ensure_ascii=False)
    flat["previous_employments"] = json.dumps(data.get("previous_employments") or [], ensure_ascii=False)
    flat["source_filename"] = filename
    flat["uploaded_at"] = datetime.now().strftime("%m/%d/%Y %I:%M:%S %p")
    return flat


def append_to_excel(row: Dict[str, Any], excel_path: str = "resumes.xlsx") -> None:
    """Append a row to the Excel file, creating it with headers if missing."""
    headers = list(RESUME_SCHEMA.keys()) + ["source_filename", "uploaded_at"]

    if os.path.exists(excel_path):
        wb = load_workbook(excel_path)
        ws = wb.active
    else:
        wb = Workbook()
        ws = wb.active
        ws.title = "Resumes"
        ws.append(headers)

    ws.append([row.get(h, "") for h in headers])
    wb.save(excel_path)


def extract_resume_data(file_content: bytes, filename: str,
                        excel_path: str = "resumes.xlsx") -> Dict[str, Any]:
    """Full pipeline: text extraction -> GPT extraction -> append to Excel."""
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".docx":
        text = extract_text_from_docx(file_content)
    elif ext == ".doc":
        # Convert legacy .doc to .docx first (requires LibreOffice)
        converted = convert_doc_to_docx(file_content)
        text = extract_text_from_docx(converted)
    elif ext in (".pdf", ".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".heif"):
        text = extract_text_with_ocr(file_content)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    if not text or len(text.strip()) < 50:
        raise ValueError("Insufficient text extracted from document")

    extracted = extract_with_gpt(text)
    row = flatten_for_excel(extracted, filename)
    append_to_excel(row, excel_path)
    return row